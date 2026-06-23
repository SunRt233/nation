from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor


ROOT = Path(r"C:\Obsidion\妙妙屋")
OUTLINE_PATH = ROOT / r"04-课件\试点产出\2026-06-17-键的极性与分子极性-基础班-上课大纲.docx"
HANDOUT_PATH = ROOT / r"06-学生侧材料\讲义\2026-06-17-键的极性与分子极性-基础班-学生讲义.docx"

CN_FONT = "Microsoft YaHei"
EN_FONT = "Times New Roman"
TITLE_COLOR = RGBColor(23, 50, 77)
SUB_COLOR = RGBColor(91, 105, 117)


def set_run_font(run, size=12, bold=False, color=None):
    run.font.name = EN_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), EN_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), EN_FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def set_paragraph_style(paragraph, spacing_after=6, spacing_before=0, line=1.3):
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(spacing_after)
    fmt.space_before = Pt(spacing_before)
    fmt.line_spacing = line


def add_text(paragraph, text, size=12, bold=False, color=None):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return run


def base_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = EN_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    normal._element.rPr.rFonts.set(qn("w:ascii"), EN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), EN_FONT)
    normal.font.size = Pt(11.5)
    return doc


def title_block(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_style(p, spacing_after=3, line=1.15)
    add_text(p, title, size=24, bold=True, color=TITLE_COLOR)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_style(p2, spacing_after=12, line=1.15)
    add_text(p2, subtitle, size=11, color=SUB_COLOR)


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_paragraph_style(p, spacing_before=8 if level == 1 else 4, spacing_after=4, line=1.2)
    add_text(p, text, size=15 if level == 1 else 12.5, bold=True, color=TITLE_COLOR)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.74 + 0.5 * level)
    p.paragraph_format.first_line_indent = Cm(-0.45)
    set_paragraph_style(p, spacing_after=3)
    add_text(p, text, size=11.5)
    return p


def para(doc, text):
    p = doc.add_paragraph()
    set_paragraph_style(p, spacing_after=5)
    add_text(p, text, size=11.5)
    return p


def simple_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = False
    widths = [Cm(4), Cm(5), Cm(7)]
    if len(rows[0]) == 4:
        widths = [Cm(3.2), Cm(3.8), Cm(6.2), Cm(4.0)]
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.width = widths[c]
            cell_p = cell.paragraphs[0]
            set_paragraph_style(cell_p, spacing_after=0, line=1.15)
            add_text(cell_p, value, size=10.5, bold=(r == 0), color=TITLE_COLOR if r == 0 else None)
    doc.add_paragraph()


def build_outline():
    doc = base_doc()
    title_block(
        doc,
        "键的极性与分子极性（基础班）上课大纲",
        "第一轮基础班桥梁增强课｜45 min｜教师课堂执行稿",
    )

    heading(doc, "一、课型定位")
    bullet(doc, "定位：分子结构基础之后的桥梁增强课，专门处理“键有极性 ≠ 分子一定有极性”这一高频混淆。")
    bullet(doc, "主链：电负性 → 键的极性 → 分子极性。")
    bullet(doc, "外扩链：VSEPR 理论 → 构型对称性 → 键矩抵消。")

    heading(doc, "二、教学目标")
    bullet(doc, "知识：分清电负性、键极性、分子极性三个层次。")
    bullet(doc, "理解：会用电负性差与构型对称性解释 CO₂、H₂O、BF₃、NH₃、CCl₄、CH₂Cl₂ 的极性。")
    bullet(doc, "应用：学生能按“先看键、再看形、最后看矢量和”的顺序完成常见主族分子的极性判断。")

    heading(doc, "三、重难点")
    simple_table(
        doc,
        [
            ["维度", "内容", "课堂处理"],
            ["重点", "键极性与分子极性的区别", "用 CO₂ / H₂O 先建立“局部 vs 整体”的对照。"],
            ["难点", "高度对称构型中的极性抵消", "至少画出 CO₂、BF₃、CCl₄ 的构型与键矩方向。"],
            ["难点", "同骨架下取代不对称导致的极性保留", "用 CCl₄ / CH₂Cl₂ / CH₃Cl 形成一组对照。"],
        ],
    )

    heading(doc, "四、课时分配")
    simple_table(
        doc,
        [
            ["时间", "板块", "教师动作", "媒介"],
            ["0-6 min", "问题切入", "先抛出 H₂O / CO₂ 与 CCl₄ 两个问题，逼学生暴露混淆。", "PPT"],
            ["6-14 min", "三层概念分家", "把电负性、键极性、分子极性三个层次讲清。", "PPT + 板书"],
            ["14-24 min", "CO₂ / H₂O 对照", "把“局部极性不等于整体极性”讲透。", "PPT + 板书"],
            ["24-34 min", "对称与不对称构型", "带 BF₃、NH₃、CCl₄、CH₂Cl₂、CH₃Cl 做分类。", "PPT"],
            ["34-41 min", "例题链", "按 HCl/HBr/HI → CO₂/H₂O → BF₃/NH₃/CCl₄/CH₃Cl 推进。", "PPT + 板书"],
            ["41-45 min", "易错点收口", "回到流程卡与易错点总表，让学生自己复述顺序。", "PPT"],
        ],
    )

    heading(doc, "五、板书框架")
    para(doc, "建议主板书固定成三行：")
    bullet(doc, "1. 电负性看原子；键极性看局部；分子极性看整体。")
    bullet(doc, "2. 判断顺序：先看键 → 再看形 → 后看矢量和。")
    bullet(doc, "3. 经典对照：CO₂ / H₂O；BF₃ / NH₃；CCl₄ / CH₂Cl₂ / CH₃Cl。")

    heading(doc, "六、教师提醒")
    bullet(doc, "不要把这节课讲成离子键专题；极性键与离子键的边界只作连续谱提醒。")
    bullet(doc, "不要把分子间作用力拉进正文，只在课末点一下“极性会影响后续物性学习”。")
    bullet(doc, "这一课的成功标准不是多讲，而是让学生以后看到新分子也能按固定顺序判断。")

    heading(doc, "七、下节衔接")
    para(doc, "可自然衔接到分子间作用力、范德华力、氢键，或者回到杂化轨道理论，把构型判断进一步连接到成键解释。")

    OUTLINE_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTLINE_PATH))


def build_handout():
    doc = base_doc()
    title_block(
        doc,
        "键的极性与分子极性（基础班）学生讲义",
        "第一轮基础班桥梁增强课｜分子结构判断专题",
    )

    heading(doc, "一、学习目标")
    bullet(doc, "能区分电负性、键的极性、分子极性三个概念。")
    bullet(doc, "能根据电负性差判断一根键是否有极性。")
    bullet(doc, "能根据构型对称性判断分子极性是否抵消。")
    bullet(doc, "能解释 CO₂、H₂O、BF₃、NH₃、CCl₄、CH₂Cl₂、CH₃Cl、SO₂、XeF₄ 等常见分子的极性。")

    heading(doc, "二、三个概念先分清")
    simple_table(
        doc,
        [
            ["概念", "研究对象", "核心问题"],
            ["电负性", "分子中的原子", "谁更会拉电子"],
            ["键的极性", "某一根化学键", "电子云偏向哪一端"],
            ["分子极性", "整个分子", "所有偏向加起来还剩不剩"],
        ],
    )
    para(doc, "一句话记忆：电负性看原子，键极性看局部，分子极性看整体。")

    heading(doc, "三、键为什么会有极性")
    para(doc, "当成键两原子的电负性不同，共用电子对会偏向电负性更大的一侧，于是形成极性键。")
    simple_table(
        doc,
        [
            ["键", "课堂判断", "说明"],
            ["Cl-Cl", "非极性键", "成键原子相同，电负性差为 0。"],
            ["H-Cl", "极性键", "共用电子对偏向 Cl。"],
            ["Na-Cl", "离子性很强", "电子转移程度很大，但不能把“极性键”和“离子键”混为一谈。"],
        ],
    )
    para(doc, "这一层先解决“某一根键偏不偏”的问题。")

    heading(doc, "四、分子极性为什么还要看构型")
    para(doc, "即使每根键都有极性，整个分子也未必有极性，因为分子极性取决于所有键矩的矢量和。")
    simple_table(
        doc,
        [
            ["分子", "构型", "分子极性", "原因"],
            ["CO₂", "直线形", "非极性", "两个 C=O 键矩大小相等、方向相反，完全抵消。"],
            ["H₂O", "V 形", "极性", "两个 O-H 键矩不能完全抵消。"],
        ],
    )
    para(doc, "这组对照一定要记住：局部极性不等于整体极性。")

    heading(doc, "五、对称构型与不对称构型")
    simple_table(
        doc,
        [
            ["分子", "构型", "分子极性", "要点"],
            ["BF₃", "平面三角形", "非极性", "对称构型中各键矩抵消。"],
            ["NH₃", "三角锥", "极性", "空间不对称，键矩有剩余。"],
            ["CCl₄", "正四面体", "非极性", "四个取代基完全相同。"],
            ["CH₂Cl₂", "四面体骨架", "极性", "取代不完全对称。"],
            ["CH₃Cl", "四面体骨架", "极性", "只有一个强极性 C-Cl 键，难以完全抵消。"],
            ["SO₂", "V 形", "极性", "与 CO₂ 对照，构型不同导致结论不同。"],
            ["XeF₄", "平面正方形", "非极性", "高度对称，键矩抵消。"],
        ],
    )
    para(doc, "判断分子极性时，除了看构型类型，还要看取代基是否相同。")

    heading(doc, "六、第一轮判断流程卡")
    bullet(doc, "1. 先写 Lewis 结构，确认中心原子周围有几组电子对。")
    bullet(doc, "2. 用 VSEPR 判断空间构型。")
    bullet(doc, "3. 判断每根键是否有极性。")
    bullet(doc, "4. 检查这些键矩能否因对称性而抵消。")

    heading(doc, "七、课堂例题")
    bullet(doc, "例 1：比较 HCl、HBr、HI 三种键的极性大小规律。")
    bullet(doc, "例 2：判断 CO₂ 和 H₂O 的分子极性，并解释原因。")
    bullet(doc, "例 3：判断 BF₃、NH₃、CCl₄、CH₃Cl 的分子极性。")
    bullet(doc, "例 4：判断 CO₂、SO₂、CH₂Cl₂、XeF₄ 的分子极性，并说明依据。")

    heading(doc, "八、核心结论")
    bullet(doc, "键的极性看电负性差，分子极性看键矩矢量和。")
    bullet(doc, "有极性键的分子不一定是极性分子。")
    bullet(doc, "构型对称性和取代对称性都会影响分子极性。")

    heading(doc, "九、易错点预警")
    simple_table(
        doc,
        [
            ["错误想法", "为什么错", "正确理解"],
            ["有极性键就一定是极性分子", "忽略构型和矢量抵消", "先看键，再看形，再看和"],
            ["极性键就是离子键", "把“偏”误解成“转移”", "极性共价键仍然是共用电子"],
            ["对称就是看上去差不多", "没有先确定空间构型", "必须先确定 VSEPR 构型"],
            ["四面体分子一定非极性", "忽略取代是否相同", "只有高度对称且取代相同才易完全抵消"],
        ],
    )

    heading(doc, "十、课后练习")
    bullet(doc, "1. 判断并解释：CO₂、SO₂、BF₃、NH₃、CCl₄、CH₂Cl₂、CH₃Cl 的分子极性。")
    bullet(doc, "2. 比较 HCl、HBr、HI 三种键的极性大小规律。")
    bullet(doc, "3. 用一句话解释为什么 CCl₄ 有极性键却是非极性分子。")
    bullet(doc, "4. 判断 PF₅ 与 XeF₄ 的分子极性，并说出关键理由。")

    heading(doc, "十一、学生自查清单")
    bullet(doc, "我能区分电负性、键极性、分子极性三个层次。")
    bullet(doc, "我能用 CO₂ / H₂O 说明“局部极性不等于整体极性”。")
    bullet(doc, "我能用 CCl₄ / CH₂Cl₂ 说明“构型相同不等于极性相同”。")
    bullet(doc, "我能按“先键、再形、后矢量和”的顺序完成常见分子极性判断。")

    HANDOUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(HANDOUT_PATH))


if __name__ == "__main__":
    build_outline()
    build_handout()
