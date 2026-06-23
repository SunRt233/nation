from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"C:\Obsidion\妙妙屋")
OUTLINE_PATH = ROOT / r"04-课件\试点产出\2026-06-17-晶体结构类型与稳定性-基础班-上课大纲.docx"
HANDOUT_PATH = ROOT / r"06-学生侧材料\讲义\2026-06-17-晶体结构类型与稳定性-基础班-学生讲义.docx"
IMG_DIR = ROOT / r"07-资料提炼\网课资料\无机化学-新课-周坤-2020-难度适中\学生讲义\04-05第三章晶体与晶体结构学生版_images"

CN_FONT = "Microsoft YaHei"
EN_FONT = "Times New Roman"
TITLE_COLOR = RGBColor(23, 50, 77)
SUB_COLOR = RGBColor(91, 105, 117)


def set_run_font(run, size=12, bold=False, color=None):
    run.font.name = EN_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), EN_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), EN_FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def set_paragraph_style(paragraph, spacing_after=6, spacing_before=0, line=1.3):
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(spacing_after)
    fmt.space_before = Pt(spacing_before)
    fmt.line_spacing = line


def add_text(paragraph, text, size=12, bold=False, color=None):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return run


def base_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = EN_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    normal._element.rPr.rFonts.set(qn("w:ascii"), EN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), EN_FONT)
    normal.font.size = Pt(11.5)
    return doc


def title_block(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_style(p, spacing_after=3, line=1.15)
    add_text(p, title, size=24, bold=True, color=TITLE_COLOR)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_style(p2, spacing_after=12, line=1.15)
    add_text(p2, subtitle, size=11, color=SUB_COLOR)


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_paragraph_style(p, spacing_before=8 if level == 1 else 4, spacing_after=4, line=1.2)
    add_text(p, text, size=15 if level == 1 else 12.5, bold=True, color=TITLE_COLOR)
    return p


def para(doc, text):
    p = doc.add_paragraph()
    set_paragraph_style(p, spacing_after=5)
    add_text(p, text, size=11.5)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.74 + 0.5 * level)
    p.paragraph_format.first_line_indent = Cm(-0.45)
    set_paragraph_style(p, spacing_after=3)
    add_text(p, text, size=11.5)
    return p


def simple_table(doc, rows, widths=None, font_size=10.5):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = False
    if widths is None:
        widths = [Cm(3.2)] * len(rows[0])
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.width = widths[c]
            cell_p = cell.paragraphs[0]
            set_paragraph_style(cell_p, spacing_after=0, line=1.15)
            add_text(cell_p, str(value), size=font_size, bold=(r == 0), color=TITLE_COLOR if r == 0 else None)
    doc.add_paragraph()


def picture_block(doc, image_name, caption, width_cm=12.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_style(p, spacing_after=2, spacing_before=2, line=1.0)
    run = p.add_run()
    run.add_picture(str(IMG_DIR / image_name), width=Cm(width_cm))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_style(cap, spacing_after=6, line=1.0)
    add_text(cap, caption, size=9.5, color=SUB_COLOR)


def build_outline():
    doc = base_doc()
    title_block(
        doc,
        "晶体结构类型与稳定性（基础班）上课大纲",
        "第一轮基础班新授课｜3 课时 × 45 min｜教师课堂执行稿",
    )

    heading(doc, "一、课型定位")
    bullet(doc, "定位：第一轮基础班晶体结构专题，用“堆积—填隙—稳定性—物性”主链把教材第 13 章主干讲实。")
    bullet(doc, "主链：等径球堆积 → 空隙与填隙 → 典型离子结构 → 稳定性规则 → 晶体类型与物性。")
    bullet(doc, "外扩链：几何规则链（空间利用率、半径比）与修正规则链（Pauling 规则、离子极化、氢键型分子晶体）。")

    heading(doc, "二、教学目标")
    bullet(doc, "知识：掌握点阵、结构基元、晶胞、密堆积、典型离子晶体结构与四类晶体的基本语言。")
    bullet(doc, "理解：能用半径比规则、Pauling 规则和离子极化解释典型结构选择。")
    bullet(doc, "应用：能从结构解释硬度、熔点、导电性、密度等典型物性。")

    heading(doc, "三、教学重难点")
    simple_table(
        doc,
        [
            ["维度", "内容", "课堂处理"],
            ["重点", "堆积模型、填隙语言、典型离子结构", "反复用“骨架 + 空隙 + 配位数”说法回收，防止只背结构名。"],
            ["重点", "四类晶体与物性关系", "用金刚石 / 石墨 / 干冰 / 冰 / 铜做桥梁例子。"],
            ["难点", "半径比与极化修正规则", "先给几何直觉，再用 AgI、CuCl 反例说明规则不是万能的。"],
            ["难点", "SiO₂ vs CO₂ 结构差异", "专门做一组“化学式相似但结构完全不同”的对照。"],
        ],
        widths=[Cm(2.4), Cm(5.0), Cm(8.6)],
    )

    heading(doc, "四、课时分配")
    simple_table(
        doc,
        [
            ["课时", "板块", "核心任务", "媒介"],
            ["课时 1", "晶体语言与等径球堆积", "讲清晶体、点阵、晶胞、sc/bcc/ccp/hcp、空隙与晶胞计数。", "PPT + 板书"],
            ["课时 2", "典型离子晶体与稳定性规则", "讲透 NaCl / CsCl / ZnS / CaF₂，补半径比、Pauling 与极化。", "PPT + 板书"],
            ["课时 3", "其他晶体类型与物性解释", "讲四类晶体，做 SiO₂ / CO₂、金刚石 / 石墨、冰 / 水等桥梁对照。", "PPT + 板书"],
        ],
        widths=[Cm(2.2), Cm(4.2), Cm(7.5), Cm(2.5)],
    )

    heading(doc, "五、板书框架")
    bullet(doc, "左侧主线固定：晶体语言 → 堆积 → 空隙 → 填隙 → 半径比 / Pauling / 极化 → 物性。")
    bullet(doc, "右侧临时区：空间利用率几何关系、半径比直觉、例题解答路径。")
    bullet(doc, "必须板书的课堂口令：NaCl=fcc+八面体全填；ZnS=fcc+四面体半填；CaF₂=fcc+四面体全填；CsCl=两套简单立方互穿。")

    heading(doc, "六、教师提醒")
    bullet(doc, "不要把这节课讲成纯晶体学符号课，晶系和点阵只保留第一轮够用口径。")
    bullet(doc, "不要只给结构图不给语言，所有典型结构都要求学生会用“骨架 + 填隙 + 配位数”复述。")
    bullet(doc, "要主动强调“几何上可行”不等于“真实稳定结构”，为后面的结构化学思维埋钩子。")
    bullet(doc, "四类晶体部分一定回到性质，不要停留在分类表背诵。")

    heading(doc, "七、下节衔接")
    para(doc, "可自然衔接到配位化学基础：本节建立了“配位数、空间构型、稳定性”的结构直觉，下节转入中心离子与配体组成的有限配位体系。")

    OUTLINE_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTLINE_PATH))


def build_handout():
    doc = base_doc()
    title_block(
        doc,
        "晶体结构类型与稳定性（基础班）学生讲义",
        "第一轮基础班新授课｜普通化学原理 + 无机结构化学桥梁版",
    )

    heading(doc, "一、学习目标")
    bullet(doc, "能用“点阵 + 结构基元 + 晶胞”的语言描述晶体。")
    bullet(doc, "能比较简单立方、bcc、ccp、hcp 的结构特征，并会做基础晶胞计数。")
    bullet(doc, "能从“填隙”视角理解 NaCl、CsCl、ZnS、CaF₂ 等典型离子结构。")
    bullet(doc, "能用半径比规则、Pauling 规则和离子极化解释结构选择。")
    bullet(doc, "能区分离子晶体、分子晶体、原子晶体、金属晶体，并解释典型物性。")

    heading(doc, "二、晶体结构的最小语言")
    para(doc, "晶体的根本特征不是“长得规整”，而是微观上具有长程有序的周期性排列。描述晶体结构时，最常用的三件套是：点阵、结构基元和晶胞。")
    para(doc, "教材里还会把晶体与非晶体放在一起比较。真正值得保留的第一轮口径有三点：晶体具有长程有序，常表现出各向异性，并有较固定的熔点；非晶体通常只有近程有序，软化范围较宽，更接近“冻结了的无序结构”。")
    simple_table(
        doc,
        [
            ["概念", "含义", "课堂提醒"],
            ["点阵", "数学上的重复骨架", "强调“在哪里重复”"],
            ["结构基元", "每个阵点真正代表的化学内容", "强调“重复的到底是什么”"],
            ["晶胞", "能平移复制整个晶体的最小重复单元", "不一定等于一个完整分子"],
        ],
        widths=[Cm(2.7), Cm(6.5), Cm(6.0)],
    )
    bullet(doc, "晶体结构 = 点阵 + 结构基元。")
    bullet(doc, "晶胞图中“看见的粒子数”不等于真正归属本晶胞的粒子数。")
    bullet(doc, "第一轮对晶系的要求是“会认框架、不过度陷入符号学”：立方、六方、四方、正交是最常见的课堂识别对象。")
    bullet(doc, "“自范性”可以知道，但不要把它当成晶体本质；本质仍然是内部周期性与长程有序。")
    simple_table(
        doc,
        [
            ["位置", "计数规则"],
            ["顶点", "每个只算 1/8"],
            ["棱心", "每个只算 1/4"],
            ["面心", "每个只算 1/2"],
            ["体心", "每个算 1"],
        ],
        widths=[Cm(4.0), Cm(10.0)],
    )
    para(doc, "因此：简单立方晶胞含 1 个粒子，bcc 含 2 个粒子，fcc 含 4 个粒子。")
    para(doc, "教材口径里，点阵回答“周期形式”，结构基元回答“化学内容”，晶胞回答“最小重复单元”。三者一旦混淆，后面所有结构判断都会失真。")
    para(doc, "再往前走半步，晶胞参数 a、b、c 与夹角 α、β、γ 决定了晶胞几何形状；这也是为什么后面谈立方、六方、四方时，实际上是在用几何参数区分不同点阵框架。")
    para(doc, "网课里这一段最值得保留的课堂口令是：不要把“规则外形”错当成本质。晶体外形可以被破坏，但只要内部周期性和长程有序还在，它依然是晶体；反过来，外形再规整，若内部没有长程周期，也不能按晶体处理。")
    picture_block(doc, "5510ab3f7cd0fa77d874b65b40850737002bea5a3d79931e8e02c30bc8986331.jpg", "图 1 14 种布拉维点阵示意。课堂只需先建立“晶系/点阵/基元是三层语言”这一意识。", width_cm=13.0)

    heading(doc, "三、等径球堆积与空隙")
    para(doc, "把原子近似看成等径小球，可以得到最常见的堆积模型。它们决定了配位数、空间利用率以及后面离子晶体的填隙方式。")
    simple_table(
        doc,
        [
            ["堆积类型", "配位数", "空间利用率", "关键几何关系"],
            ["简单立方", "6", "约 52%", "a = 2r"],
            ["bcc", "8", "约 68%", "4r = sqrt(3)a"],
            ["ccp / fcc", "12", "约 74%", "4r = sqrt(2)a"],
            ["hcp", "12", "约 74%", "与 ccp 同为最密堆积"],
        ],
        widths=[Cm(3.6), Cm(2.0), Cm(3.0), Cm(6.0)],
    )
    bullet(doc, "ccp 与 hcp 的差别主要在堆叠顺序：ABCABC... vs ABAB...")
    bullet(doc, "密堆积中每个粒子平均对应 1 个八面体空隙和 2 个四面体空隙。")
    bullet(doc, "几何公式要会回收：a = 2r（简单立方），4r = sqrt(3)a（bcc），4r = sqrt(2)a（fcc）。")
    bullet(doc, "bcc 虽然不是最密堆积，但不能简单说它“不稳定”；它在很多金属中依然是稳定结构，说明真实稳定性不只由空间利用率单独决定。")
    para(doc, "提醒：最密堆积只是几何上很紧，不自动等于一切真实体系都最稳定。")
    para(doc, "普通化学原理这一章真正想建立的，不只是几个百分数，而是“堆积方式会决定配位数、空隙类型和后续可形成的晶体结构”。")
    para(doc, "从解题角度看，等径球堆积最常回收的有三类任务：由 a 求 r，由结构求空间利用率，由密堆积求四面体/八面体空隙数。学生至少要把这三类回收动作建立起来。")
    para(doc, "若把这部分再说得更课堂化一点，可以记成一句话：堆积模型本身不是考点终点，而是后面一切“结构像什么、空隙怎么占、配位为什么这样、性质为何不同”的几何出发点。")
    picture_block(doc, "b0b76255925b5a3bb6fd74597da2c9bffa748735e32831ccfcc8c92e4aec83e4.jpg", "图 2 六方最密堆积的层状直观图。抓住 ABAB 堆叠和层间空隙的来源。", width_cm=13.0)

    heading(doc, "四、典型离子晶体：从“填隙”看结构")
    para(doc, "离子晶体常可理解为“大离子先堆积，小离子再填隙”。判断典型结构时，要养成“先看骨架，再看空隙，再看配位数”的习惯。")
    simple_table(
        doc,
        [
            ["结构类型", "骨架", "填隙方式", "配位数", "晶胞内化学式单位"],
            ["NaCl 型", "阴离子 fcc", "全部八面体空隙", "6:6", "4"],
            ["CsCl 型", "两套简单立方互穿", "立方体中心位", "8:8", "1"],
            ["立方 ZnS 型", "阴离子 fcc", "一半四面体空隙", "4:4", "4"],
            ["CaF₂ 型", "Ca2+ fcc", "全部四面体空隙", "8:4", "4"],
        ],
        widths=[Cm(3.2), Cm(3.8), Cm(4.3), Cm(2.0), Cm(2.7)],
    )
    bullet(doc, "课堂口令：NaCl = fcc + 八面体全填；ZnS = fcc + 四面体半填；CaF₂ = fcc + 四面体全填。")
    bullet(doc, "CsCl 的典型误区：外形“像 bcc”，但顶点和体心不是同种粒子，所以不是金属那种体心立方。")
    para(doc, "从 AB 型走向 AB₂ 型，说明晶体结构不是只背单一模板，而是一个“骨架 + 填隙 + 配位 + 化学计量比”的统一语言。")
    para(doc, "进一步说，NaCl、ZnS、CaF₂ 三者都和 fcc 骨架相关，但因为填入空隙的种类与占据率不同，配位关系、化学式单位数和物性图景都跟着变化。")
    para(doc, "教材和无机化学里还常顺手提到 TiO₂（金红石型）与钙钛矿型结构。第一轮不要求完全掌握所有细节，但要知道：当化学计量比从 AB 扩展到 AB₂、ABX₃ 时，填隙语言仍然有效，只是骨架和占据空隙的对象更复杂。")
    bullet(doc, "AB₂ 的第一层认识：CaF₂ 可看作 fcc 骨架上四面体空隙全填；TiO₂ 则常回收到“配位数不同、八面体环境不同”的比较。")
    bullet(doc, "ABX₃ 的第一层认识：钙钛矿型强调“大离子定骨架，小而高价离子进八面体空隙”，常用来解释 BaTiO₃、SrTiO₃ 这类材料。")
    para(doc, "如果把“化学计量比变化”也纳入结构判断，就能更自然地理解为什么不能只问“它像不像 NaCl”。真正要问的是：骨架是谁、空隙是谁占、占了多少、这样得到的配位数与计量比是否同时说得通。")
    para(doc, "这一段也是教材与网课共同强调的结构化学意识：化学式只告诉我们“比例”，真正决定晶体类型的是比例如何在空间里实现。换句话说，化学计量比是约束条件，不是结构答案本身。")
    picture_block(doc, "cc1e9606a12e4230a10339e61d8f6d7090f87672437a6b1741c5414d57751a4e.jpg", "图 3 典型立方结构示意。课上可用来对照 CsCl“像 bcc 但不是 bcc”的误区。", width_cm=13.0)

    heading(doc, "五、为什么不是“能排就行”：稳定性规则")
    para(doc, "真实晶体结构既要几何上放得下，也要静电上划算，还要考虑键型是否已经发生偏离。第一轮够用的三条规则如下。")
    simple_table(
        doc,
        [
            ["规则", "主要回答的问题", "课堂直觉"],
            ["半径比规则", "几配位在几何上可行", "能不能坐进去"],
            ["Pauling 规则", "静电稳定性是否合理", "静电账算不算得过来"],
            ["离子极化", "为什么会偏离理想离子模型", "它是不是开始不像纯离子晶体了"],
        ],
        widths=[Cm(3.0), Cm(6.0), Cm(5.3)],
    )
    para(doc, "半径比的基础口径：较小阳离子倾向 4 配位，中等倾向 6 配位，更大倾向 8 配位。")
    para(doc, "若把配位数看成“阳离子周围阴离子的拥挤程度”，那么半径比规则本质上是在问：这个阳离子在不让同号阴离子彼此过度靠近的前提下，最多能稳稳占据什么环境。")
    para(doc, "Pauling 规则在基础班不必展开成复杂电价计算，但至少要知道：阴离子周围每条静电键分摊的稳定作用需要‘加总得上’，否则几何上能塞进去，静电上也可能不划算。")
    para(doc, "第一轮够用的 Fajans 直觉：阳离子越小、价态越高，极化力通常越强；阴离子越大、越易被极化，共价性越容易增强，于是结构可能偏向低配位、方向性更强的形式。")
    bullet(doc, "典型修正例子：AgI、CuCl 等体系不能只靠理想半径比判断。")
    bullet(doc, "课堂答题模板：先给几何判断，再补静电稳定性，最后看是否有极化导致的偏离。")
    bullet(doc, "要会分清两句话：极化力强，说的是阳离子“拉扯电子云”的能力；易极化性强，说的是阴离子“容易被拉变形”的程度。")
    para(doc, "若用教材风格的表达来概括这一段，可以记成“三层筛选”：半径比看几何能否成立，Pauling 规则看静电账是否平，极化效应看键型是否已开始偏离纯离子模型。这样答题会比只背阈值更稳定。")

    heading(doc, "六、四类晶体与结构解释物性")
    simple_table(
        doc,
        [
            ["晶体类型", "结构单元", "主要作用", "典型性质", "代表物"],
            ["离子晶体", "离子", "离子键", "熔点高、硬脆、固态不导电", "NaCl"],
            ["分子晶体", "分子", "分子间力 / 氢键", "熔点较低、较软、不导电", "干冰、冰、I2"],
            ["原子晶体", "原子", "无限共价网络", "熔点极高、很硬、多数不导电", "金刚石、SiO2"],
            ["金属晶体", "金属原子骨架", "金属键", "导电导热、具延展性", "Cu、Fe"],
        ],
        widths=[Cm(2.8), Cm(2.5), Cm(3.6), Cm(5.0), Cm(2.3)],
    )
    para(doc, "四类晶体的判断最后一定要落到性质。考试里最有区分度的，不是能不能背出分类表，而是能不能说清“为什么它这么硬 / 这么低熔 / 能导电 / 密度反常”。")
    para(doc, "原子晶体部分建议额外记住两个桥梁对象：SiO₂ 用来对照 CO₂，石墨用来提醒“同为碳，层内层间作用却完全不同”。这两个例子是普通化学原理和无机化学里都很高频的结构解释题。")
    para(doc, "分子晶体也不要只记“低熔点”。还要会分层：近球形分子常能取较紧堆积，线形或定向性明显的分子则会受取向影响；冰这种氢键分子晶体更要强调“方向性强但堆积不紧”，这正是密度反常的根源。")
    para(doc, "金属晶体部分则应多说一句：金属键无明显方向性，所以金属更容易形成高配位、较紧密堆积，也因此常兼具导电、导热与延展性。")
    para(doc, "原子晶体的答题口径最好固定为：‘整个晶体就是一个巨大的共价网络’，这样就不会误把它理解成“分子之间作用很强”的分子晶体。")
    para(doc, "教材里与竞赛衔接最有价值的，不是机械背四类表格，而是反复追问：结构单元是谁、粒子之间靠什么结合、这种结合为什么会导向特定物性。只要这三句问对，绝大多数判断题与解释题都会更稳定。")
    para(doc, "如果要把这一节真正学“活”，还要注意一个常见迁移：分类表只是入口，最终落点是解释。也就是说，考试里更高分的答案，往往不是“它属于某某晶体”，而是“因为它的结构单元与作用方式如此，所以表现出某某性质”。")

    heading(doc, "七、几个必须会说的桥梁对照")
    bullet(doc, "金刚石 vs 石墨：同为碳，但前者是三维共价网络，后者层内强、层间弱，所以硬度与导电性完全不同。")
    bullet(doc, "冰 vs 液态水：冰中水分子靠方向性很强的氢键形成开放网络，空间利用率低，所以密度更小。")
    bullet(doc, "SiO₂ vs CO₂：化学计量比相似，但 SiO₂ 是三维共价骨架原子晶体，CO₂ 是有限分子形成的分子晶体。")
    bullet(doc, "NaCl vs ZnS：都可看作阴离子 fcc 骨架，但分别填八面体与四面体空隙，因此配位数不同。")
    bullet(doc, "C₆₀ vs 金刚石：前者是分子晶体，球形分子之间主要靠较弱作用；后者是原子晶体，整个晶体被共价键贯通。")
    bullet(doc, "AgI vs NaCl：都可写成 AB 型，但前者更容易体现极化导致的结构偏离，后者更接近典型离子模型。")

    heading(doc, "八、当堂应解决的核心追问")
    bullet(doc, "为什么 CsCl 不能简单叫作 bcc？关键不在外形像不像，而在顶点和体心是否属于同种粒子。")
    bullet(doc, "为什么 NaCl 和 ZnS 都能从阴离子 fcc 骨架出发，却得出不同结构？关键在填入的是八面体还是四面体空隙。")
    bullet(doc, "为什么半径比能给出配位倾向，却不是最后裁判？因为真实结构还要过静电稳定性和极化修正两关。")
    bullet(doc, "为什么冰、石墨、SiO₂ 这类题总爱考？因为它们最能训练“结构解释性质”的能力，而不是死背分类名词。")

    heading(doc, "九、题目里最常用的解题动作")
    simple_table(
        doc,
        [
            ["题型", "先做什么", "再做什么"],
            ["结构识别题", "辨骨架与填隙", "数配位数、核对化学式单位"],
            ["物性解释题", "先说结构单元与作用力", "再落到熔点/硬度/导电性"],
            ["规则判断题", "先给半径比直觉", "再补 Pauling / 极化修正"],
            ["对比题", "先找共同点", "再抓决定差异的那一层结构"],
        ],
        widths=[Cm(3.0), Cm(5.5), Cm(7.0)],
    )
    para(doc, "把这张表背后的顺序真正练熟，很多题就不会写成“想到哪写到哪”。晶体结构题最怕的不是不会，而是语言顺序混乱，导致明明知道结构却拿不到分。")

    heading(doc, "十、课堂速记口令")
    bullet(doc, "晶体结构三件套：点阵、结构基元、晶胞。")
    bullet(doc, "结构判断三连问：骨架是谁？填了什么空隙？配位数多少？")
    bullet(doc, "稳定性三层筛选：几何可行、静电合理、极化修正。")
    bullet(doc, "物性解释三步式：结构单元 → 作用方式 → 性质表现。")

    heading(doc, "十一、课堂例题")
    bullet(doc, "例 1：判断“晶胞”“结构基元”“晶体定义”相关说法的正误并说明理由。")
    bullet(doc, "例 2：比较 bcc、ccp、hcp 的配位数和空间利用率。")
    bullet(doc, "例 3：用“骨架 + 填隙 + 配位数”描述 NaCl、CsCl、ZnS。")
    bullet(doc, "例 4：解释金刚石 / 石墨 / 冰的典型物性。")
    bullet(doc, "例 5：说明为什么半径比规则不是万能的。")
    bullet(doc, "例 6：比较 SiO₂ 与 CO₂ 的晶体类型与物性差异。")
    bullet(doc, "例 7：已知某结构为 fcc 骨架且四面体空隙全填，反推配位关系与化学计量比。")
    bullet(doc, "例 8：比较 CaF₂ 与 ZnS，说明“都与四面体空隙有关”但为什么计量比与配位关系不同。")
    bullet(doc, "例 9：给出一种物质的主要物性，反推它更可能是哪类晶体，并说明理由。")

    heading(doc, "十二、核心结论")
    bullet(doc, "1. 晶体结构的最小语言是：点阵、结构基元、晶胞。")
    bullet(doc, "2. 典型结构判断主线是：先看堆积，再看填隙，再看配位数。")
    bullet(doc, "3. 真实稳定结构要同时考虑几何、静电和键型修正。")
    bullet(doc, "4. 学晶体结构的最终目的，是用结构解释性质。")
    bullet(doc, "5. 化学式相近、元素相同都不保证结构相近，真正决定性质的是空间骨架与作用方式。")

    heading(doc, "十三、易错点预警")
    simple_table(
        doc,
        [
            ["错误想法", "为什么错", "正确理解"],
            ["晶胞就是一个完整分子", "把最小重复单元和单个粒子混了", "晶胞只负责重复复制整个结构"],
            ["CsCl 就是 bcc", "忽略顶点和体心粒子种类不同", "CsCl 是两套简单立方互穿"],
            ["半径比规则是万能的", "忽略 Pauling 规则和极化修正", "结构判断要综合三层判据"],
            ["分子晶体熔点只看分子量", "忽略分子间作用力和堆积紧密程度", "还要看氢键、极性和空间利用率"],
            ["SiO2 与 CO2 化学式相似，所以结构也相似", "组成比例相似不等于空间结构相似", "要先判断是否形成无限网络"],
            ["最密堆积一定最稳定", "把几何紧密程度误当成唯一判据", "真实稳定性还受电子结构和键型影响"],
        ],
        widths=[Cm(4.0), Cm(5.0), Cm(6.0)],
    )

    heading(doc, "十四、课后作业（精简版）")
    bullet(doc, "1. 用“点阵 + 结构基元 + 晶胞”的语言解释晶体结构，并补充晶胞计数规则。")
    bullet(doc, "2. 用“骨架 + 填隙 + 配位数”的格式比较 NaCl、CsCl、ZnS、CaF₂。")
    bullet(doc, "3. 用“结构 → 作用力 → 性质”的格式比较冰、干冰、金刚石、石墨或铜中的任意三种。")
    para(doc, "说明：本专题更强调课堂内把主干框架、典型结构和性质解释讲清，因此课后作业只保留少量代表题，不再靠题量补知识缺口。")

    heading(doc, "十五、课堂例题标准答案版")
    bullet(doc, "例 1 参考要点：晶胞不是一个完整分子；晶体本质是长程有序；晶体结构可表示为点阵 + 结构基元。")
    bullet(doc, "例 2 参考要点：bcc 配位数 8、空间利用率约 68%；ccp 与 hcp 配位数都为 12、空间利用率约 74%，且二者同属最密堆积。")
    bullet(doc, "例 3 参考要点：NaCl = 阴离子 fcc 骨架 + 八面体空隙全填 + 6:6；CsCl = 两套简单立方互穿 + 8:8；ZnS = 阴离子 fcc 骨架 + 四面体空隙半填 + 4:4。")
    bullet(doc, "例 4 参考要点：金刚石是三维共价网络，所以极硬且不导电；石墨层内强层间弱，所以较软且可导电；冰因氢键开放网络而密度低于液态水。")
    bullet(doc, "例 5 参考要点：半径比只解决几何可行性；真实结构还要结合 Pauling 规则与极化修正，因此不能把半径比当万能判据。")
    bullet(doc, "例 6 参考要点：CO₂ 是有限分子形成分子晶体，SiO₂ 是三维共价骨架原子晶体，所以两者物性差异巨大。")
    bullet(doc, "例 7 参考要点：fcc 骨架若四面体空隙全填，常对应 1:2 计量关系；若是正离子作骨架、负离子全填四面体空隙，可回收到 CaF₂ 型。")
    bullet(doc, "例 8 参考要点：ZnS 与 CaF₂ 都涉及四面体空隙，但前者只填一半、后者全部填满，所以化学计量比和配位关系不同。")
    bullet(doc, "例 9 参考要点：先从熔点、硬度、导电性判断结构单元和作用力，再反推更可能属于离子晶体、分子晶体、原子晶体或金属晶体。")
    para(doc, "课堂答题提醒：标准答案最好都尽量写成“结构 / 骨架 → 作用方式或填隙方式 → 配位数或性质”这一顺序，避免只报结论不解释原因。")

    heading(doc, "十六、学生自查清单")
    bullet(doc, "我能用“点阵 + 结构基元 + 晶胞”描述晶体。")
    bullet(doc, "我能区分 bcc、ccp、hcp，并能说出它们的配位数和紧密程度。")
    bullet(doc, "我能从“填隙”角度讲清 NaCl、CsCl、ZnS、CaF2 的主要区别。")
    bullet(doc, "我能说出半径比规则、Pauling 规则和极化修正分别解决什么问题。")
    bullet(doc, "我能用结构解释金刚石、石墨、冰、干冰、SiO2、CO2 的典型性质。")
    bullet(doc, "我知道“最密堆积”“高配位”“真实稳定结构”不是完全等价的三个概念。")
    bullet(doc, "我能用较完整的一句话回答“它为什么属于这种晶体、为什么有这种性质”。")

    OUTLINE_PATH.parent.mkdir(parents=True, exist_ok=True)
    HANDOUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(HANDOUT_PATH))


if __name__ == "__main__":
    build_outline()
    build_handout()
